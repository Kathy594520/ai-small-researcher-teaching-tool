from pathlib import Path
from urllib.request import Request, urlopen
import json
import os
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


ROWS = [
    {
        "file": "apatw_taiwan_serow.jpg",
        "name": "?啁?控蝢?,
        "page": "https://www.apatw.org/project-article/9787",
        "url": "https://www.apatw.org/sites/default/files/imce/user/user46227/01_607.jpg",
        "use": "??撱??滓撅勗銋喲?",
    },
    {
        "file": "apatw_yellow_throated_marten.jpg",
        "name": "暺?鞎?,
        "page": "https://www.apatw.org/project-article/9787",
        "url": "https://www.apatw.org/sites/default/files/imce/user/user46227/02_593.jpg",
        "use": "璉格??楠?漯瘚偌撗豢ㄡ??,
    },
    {
        "file": "apatw_muntjac.jpg",
        "name": "撅梁?鈭?",
        "page": "https://www.apatw.org/project-article/9787",
        "url": "https://www.apatw.org/sites/default/files/imce/user/user46227/03_436.jpg",
        "use": "?車鈭????箄?撖?,
    },
    {
        "file": "tvbs_fishing_mongoose.png",
        "name": "?輸?撅梢??寧",
        "page": "https://news.tvbs.com.tw/life/1560301",
        "url": "https://cc.tvbs.com.tw/img/upload/2017/11/09/20171109193947-11b8e8d3.png",
        "use": "蝝?蝺璈皜祉內蝭?,
    },
    {
        "file": "river_leopard_cat.jpg",
        "name": "瞈偌皞芰??,
        "page": "https://river.udn.com/river/story/124359/8843270",
        "url": "https://uc.udn.com.tw/photo/2025/07/01/realtime/32424528.jpg",
        "use": "?唾??曇馱?眾撌???,
    },
    {
        "file": "udn_black_bear_camera_setup.jpg",
        "name": "?唳暺??嗉身蝝?蝺璈?,
        "page": "https://udn.com/news/story/7470/8693566",
        "url": "https://uc.udn.com.tw/photo/2025/04/23/realtime/31913543.jpg",
        "use": "?豢??嗉身??憭皜祉內??,
    },
]


def main():
    base = Path(__file__).resolve().parents[2]
    target = Path(os.environ["TARGET_DIR"])
    target.mkdir(exist_ok=True)
    source = base / "shallow_camera_photos"

    for name in [
        "apatw_taiwan_serow.jpg",
        "apatw_yellow_throated_marten.jpg",
        "apatw_muntjac.jpg",
        "tvbs_fishing_mongoose.png",
    ]:
        src = source / name
        if src.exists():
            shutil.copy2(src, target / name)

    headers = {"User-Agent": "Mozilla/5.0"}
    for item in ROWS:
        req = Request(item["url"], headers=headers)
        with urlopen(req, timeout=40) as r:
            data = r.read()
        (target / item["file"]).write_bytes(data)

    md_lines = [
        "# 瘛箏控?芸??豢??抒?靘?銵?,
        "",
        "| 瑼? | ?抒??批捆 | 靘?? | ???湧? | ?飛?券?|",
        "|---|---|---|---|---|",
    ]
    for row in ROWS:
        md_lines.append(
            f"| {row['file']} | {row['name']} | {row['page']} | {row['url']} | {row['use']} |"
        )
    (target / "source_table.md").write_text("\n".join(md_lines), encoding="utf-8")

    doc = Document()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("瘛箏控?芸??豢??抒?靘?銵?)
    r.bold = True
    r.font.size = Pt(16)
    doc.add_paragraph("隞乩????舐?潛?控蝢?撅梁?????蝝?蝺璈皜祆?摮詻?)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, ["瑼?", "?抒??批捆", "靘??", "???湧?", "?飛?券?]):
        cell.text = text
    for row in ROWS:
        cells = table.add_row().cells
        cells[0].text = row["file"]
        cells[1].text = row["name"]
        cells[2].text = row["page"]
        cells[3].text = row["url"]
        cells[4].text = row["use"]
    doc.save(str(target / "source_table.docx"))

    manifest = []
    for f in sorted(target.iterdir()):
        if f.is_file() and f.suffix.lower() in {".jpg", ".png"}:
            manifest.append({"file": f.name, "size": f.stat().st_size})
    (target / "manifest.json").write_text(
        json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8"
    )


if __name__ == "__main__":
    main()

