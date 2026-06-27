from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


PROJECT_ROOT = Path(__file__).resolve().parents[2]
BASE = PROJECT_ROOT
SOURCE = next(BASE.glob("AI*??.docx"))
OUTPUT = BASE / "AI撠??弦?︵?嗆?瘛箏控蝎暸?_??_?怠???docx"
MATERIALS = BASE / "materials"

CORE_IMAGES = [
    (
        "?唾?撟潮?餈",
        "Prionailurus_bengalensis_in_Taiwan_03_docx.jpg",
        "Wikimedia Commons",
    ),
    (
        "?唾???",
        "Leopard_Cat_Prionailurus_bengalensis_borneoensis_7136074909_docx.jpg",
        "Wikimedia Commons",
    ),
    (
        "摰嗉?撠",
        "Domestic_cat_felis_catus_stare_docx.jpg",
        "Wikimedia Commons",
    ),
    (
        "?豢??琿蝷箸?",
        "Wildlife_Technician_Andrew_checks_trail_cameras.jpg",
        "Wikimedia Commons",
    ),
]

SUPPLEMENTAL_IMAGES = [
    (
        "璉脣?楝",
        "Woods_landscape_road_loyalsock_state_forest_163703_docx.jpg",
        "Wikimedia Commons",
    ),
    (
        "?蝛輯??楝",
        "Bear_family_crossing_road.jpg",
        "Wikimedia Commons",
    ),
    (
        "????豢??",
        "pixy_trail_camera_5566971.jpg",
        "Pixy",
    ),
    (
        "???蝛輯?璅內",
        "pixy_wildlife_crossing_6379990.jpg",
        "Pixy",
    ),
]


def set_run_font(run, font_name="Arial", size=None, bold=None):
    run.font.name = font_name
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    rfonts = run._element.rPr.rFonts
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:eastAsia"), "Microsoft JhengHei")


def add_image_block(doc, title, filename, source_label):
    img_path = MATERIALS / filename
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(img_path), width=Inches(5.4))

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(3)
    cap.paragraph_format.space_after = Pt(12)
    cap_run = cap.add_run(f"{title}?鞈?靘?嚗source_label}")
    set_run_font(cap_run, size=9)
    cap_run.italic = True


def add_section(doc, heading, items):
    doc.add_paragraph().add_run("")
    h = doc.add_paragraph(heading, style="Heading 1")
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after = Pt(8)
    note = doc.add_paragraph()
    note.paragraph_format.space_after = Pt(6)
    note_run = note.add_run("隞乩??抒??舐?亦?潸玨?陛?晞飛??撖??飛蝷箇???)
    set_run_font(note_run, size=10)
    for title, filename, source in items:
        add_image_block(doc, title, filename, source)


def main():
    doc = Document(str(SOURCE))

    # Append a clean appendix section at the end of the lesson plan.
    doc.add_page_break()
    add_section(doc, "????????唾?颲刻??璈?曹蜓??, CORE_IMAGES)
    add_section(doc, "?????????璉脣憡????脰?????", SUPPLEMENTAL_IMAGES)

    doc.save(str(OUTPUT))
    print(f"saved {OUTPUT}")


if __name__ == "__main__":
    main()

