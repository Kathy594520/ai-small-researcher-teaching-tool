from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

SKILL_SCRIPTS = Path(r"C:\Users\USER\.codex\plugins\cache\openai-primary-runtime\documents\26.619.11828\skills\documents\scripts")
if str(SKILL_SCRIPTS) not in sys.path:
    sys.path.append(str(SKILL_SCRIPTS))

from table_geometry import apply_table_geometry

PROJECT_ROOT = Path(__file__).resolve().parents[2]


OUT_PATH = PROJECT_ROOT / "??????_???????.docx"


def set_run_font(run, *, name="Calibri", size_pt=11, bold=False, color="000000"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def format_paragraph(paragraph, *, before=0, after=6, line=1.15, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    paragraph.alignment = alignment
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = line


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def add_hyperlink(paragraph, text: str, url: str, *, color="0563C1", size_pt=10):
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_pr.append(r_fonts)
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    r_pr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    r_pr.append(u)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size_pt * 2)))
    r_pr.append(sz)
    run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def set_cell_text(cell, text: str, *, size_pt=10, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    format_paragraph(p, before=0, after=0, line=1.0)
    r = p.add_run(text)
    set_run_font(r, size_pt=size_pt, bold=bold)


def build_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.15

    for name, size_pt, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size_pt)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    title = doc.add_paragraph()
    format_paragraph(title, before=0, after=3, line=1.0)
    t_run = title.add_run("?啁?唾?靽嚗?摮豢?蝷箇蝎暹???????)
    set_run_font(t_run, size_pt=18, bold=True, color="000000")

    subtitle = doc.add_paragraph()
    format_paragraph(subtitle, before=0, after=8, line=1.0)
    s_run = subtitle.add_run("?券?蝯血飛?? NotebookLM 鞈??”?絲?望?蝪∪???箇??歇?⊥?????/ ?寞?蝔?/ 鈭車?閰?)
    set_run_font(s_run, size_pt=10.5, color="555555")

    intro = doc.add_paragraph()
    format_paragraph(intro, before=0, after=8, line=1.15)
    intro.add_run("?詨??斗?嚗?).bold = True
    intro.add_run("?祇??飛???臭誑撖怒??????唾?????????蝘蝔桃??銝遣霅啁?亙神????車?閬?鈭車嚗?閮餅??靘????鳴??踹?????憿榆?啁?隤芣?撖急??桐?摰???)

    h1 = doc.add_paragraph("銝??摰?牧瘜?)
    format_paragraph(h1, before=18, after=8, line=1.0)
    for r in h1.runs:
        set_run_font(r, size_pt=16, color="2E74B5")

    p1 = doc.add_paragraph()
    format_paragraph(p1, before=0, after=6, line=1.15)
    p1.add_run("?唾?嚗rionailurus bengalensis嚗?啁?臭??曉?????蝘蝔柴?)
    p1.add_run(" ??臬??函?冽?摮豢?蝷箝?)

    p2 = doc.add_paragraph()
    format_paragraph(p2, before=0, after=6, line=1.15)
    p2.add_run("?亥?撘瑁矽?啁?靘?嚗撖恬?")
    p2.add_run("???????唾???)
    p2.add_run(" ??")
    p2.add_run("??????蝘蝔桃??)

    p3 = doc.add_paragraph()
    format_paragraph(p3, before=0, after=8, line=1.15)
    p3.add_run("?亥??????剛降嚗遣霅啣神嚗?)
    p3.add_run("????曹??黎?其???憿?蝛嗡葉????摰??祇??飛銝??湔??銝鈭車?迂撖急?摰???)

    h2 = doc.add_paragraph("鈭迤蝣箇閰???)
    format_paragraph(h2, before=12, after=6, line=1.0)
    for r in h2.runs:
        set_run_font(r, size_pt=13, color="2E74B5")

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    for cell, txt in zip(table.rows[0].cells, ["撱箄降撖急?", "?踹?撖急?"]):
        cell.text = txt
        shade_cell(cell, "E8EEF5")
        set_cell_text(cell, txt, size_pt=10, bold=True)

    rows = [
        ("?啁?????, "?啁?寞?蝔桃??),
        ("?啁??鞎??拍車?唾?", "?啁?臭??寞?鞎??拍車"),
        ("?亙??冽??鳴???鈭車??蝢文???, "?芾酉??皞停?湔撖怨??蝔?),
        ("?唾??典???靽?暹?", "?唾??典??歇摰蝛拙???????),
    ]
    for left, right in rows:
        row = table.add_row()
        set_cell_text(row.cells[0], left, size_pt=10)
        set_cell_text(row.cells[1], right, size_pt=10)

    apply_table_geometry(table, [4680, 4680], table_width_dxa=9360, indent_dxa=120)

    h2b = doc.add_paragraph("銝?湔?暸脫?摮豢?蝷箇???")
    format_paragraph(h2b, before=14, after=6, line=1.0)
    for r in h2b.runs:
        set_run_font(r, size_pt=13, color="2E74B5")

    text_blocks = [
        "?唾??臬??銝?曉?????蝘蝔殷?銋瘛箏控??蝟餌???隞?”?銝??嫣???暑???虜銵函內?嗅????摰?鞈芰?璉脣????隞嗚?,
        "?唾?????????璉脣?渡???頝航楝畾箝?梯炊??鈭箇銵????????唾??暑?征??撠香鈭⊿◢?芸?擃?銋?憓??犖憿?瘣餌憓??拇??,
        "靽?唾???暺?銝?舫?摰喟?嚗閬???????ㄡ?啜??函?蝘餃?頝舐???頛???芷?梧?隞亙??賢?鈭粹??勗??憓恣?撘?,
        "?迨嚗???脖??臬銝?拍車霅圈?嚗?啁瘛箏控?啣????拙?璅?找?摮???撌乩???,
    ]
    for block in text_blocks:
        p = doc.add_paragraph()
        format_paragraph(p, before=0, after=6, line=1.15)
        p.add_run(block)

    h2c = doc.add_paragraph("?飛??鞈??”?????")
    format_paragraph(h2c, before=12, after=6, line=1.0)
    for r in h2c.runs:
        set_run_font(r, size_pt=13, color="2E74B5")

    key_points = [
        ("銝餉?", "?啁?唾?"),
        ("摰?", "?啁??鞎??拍車嚗?撖急??啁?寞?蝔?),
        ("憡?", "璉脣?渡???頝航楝畾箝?梯炊?犖?貉?蝒?),
        ("靽?寞?", "靽風璉脣??撠楝畾箝宏?文?芷?晞???噙雿?撘瑟???),
        ("蝯?", "靽風?唾?嚗停?臬?霅瑕??滓撅梁???),
    ]
    for label, value in key_points:
        p = doc.add_paragraph()
        format_paragraph(p, before=0, after=3, line=1.15)
        p.add_run(f"{label}嚗?).bold = True
        p.add_run(value)

    h2d = doc.add_paragraph("鈭撘靘?")
    format_paragraph(h2d, before=14, after=6, line=1.0)
    for r in h2d.runs:
        set_run_font(r, size_pt=13, color="2E74B5")

    source_table = doc.add_table(rows=1, cols=3)
    source_table.style = "Table Grid"
    for cell, txt in zip(source_table.rows[0].cells, ["靘?", "?舀????", "???"]):
        cell.text = txt
        shade_cell(cell, "E8EEF5")
        set_cell_text(cell, txt, size_pt=10, bold=True)

    sources = [
        ("?啁?唾?靽??嚗???剖?瑼?", "?啁?臭??曉?????蝘蝔柴?撣??扯?蝜?", "https://www.twlcat.org/%E8%AA%8D%E8%AD%98%E7%9F%B3%E8%99%8E/"),
        ("?啁?唾?靽??嚗???剖璈?, "靽摰??銝剔?銝餉?憡??璈瑽?, "https://www.twlcat.org/%E7%9F%B3%E8%99%8E%E7%9A%84%E5%8D%B1%E6%A9%9F/"),
        ("Vigne et al., 2016, PLOS ONE", "?唾??葉??振鞎絲皞???弦", "https://doi.org/10.1371/journal.pone.0147295"),
        ("Tamada et al., 2008, Zoological Science", "鈭散鞊寡???摮?唾??啁???", "https://doi.org/10.2108/zsj.25.154"),
        ("Chua et al., 2016, Mammal Research", "?唾??黎撖漲??蝛箏?刻?憌?, "https://doi.org/10.1007/s13364-015-0259-4"),
        ("Grassman et al., 2005, Journal of Zoology", "?唾?瘣餃?蝭????抒?蝛?, "https://doi.org/10.1017/S095283690500659X"),
        ("Johnson et al., 2006, Science", "?曆誨鞎?瞍???憿???, "https://doi.org/10.1126/science.1122277"),
    ]

    for name, note, url in sources:
        row = source_table.add_row()
        set_cell_text(row.cells[0], name, size_pt=9.5)
        set_cell_text(row.cells[1], note, size_pt=9.5)
        row.cells[2].text = ""
        p = row.cells[2].paragraphs[0]
        format_paragraph(p, before=0, after=0, line=1.0)
        add_hyperlink(p, "?????", url, size_pt=9.5)

    apply_table_geometry(source_table, [2400, 4620, 2340], table_width_dxa=9360, indent_dxa=120)

    closing = doc.add_paragraph()
    format_paragraph(closing, before=10, after=0, line=1.15)
    closing.add_run("?祇??飛撱箄降嚗?).bold = True
    closing.add_run("?冽?憿????乩葉?芸?雿輻???????芣??冽?蝣箏??函摰?憿?蝛嗆?嚗??鈭車??蝢文???)

    return doc


def main():
    doc = build_document()
    doc.save(OUT_PATH)
    print(f"saved {OUT_PATH.resolve()}")


if __name__ == "__main__":
    main()

