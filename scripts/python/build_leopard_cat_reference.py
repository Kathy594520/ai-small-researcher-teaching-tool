from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

SKILL_SCRIPTS = Path(r"C:\Users\USER\.codex\plugins\cache\openai-primary-runtime\documents\26.619.11828\skills\documents\scripts")
if str(SKILL_SCRIPTS) not in sys.path:
    sys.path.append(str(SKILL_SCRIPTS))

from table_geometry import apply_table_geometry

PROJECT_ROOT = Path(__file__).resolve().parents[2]


OUT_PATH = PROJECT_ROOT / "??????????.docx"


WEB_RESOURCES = [
    (
        "?啁?唾?靽??",
        "?典靽??嚗?????唳??脫撱?冗????ㄡ?啣?霅瑁?閮?,
        "https://www.twlcat.org/",
    ),
    (
        "颲脫平?函??拙?璅?抒?蝛嗆?",
        "?舀?唾?隤踵??蝛嗚皜祈?靽?典誨?賊??批捆??,
        "https://www.tesri.gov.tw/",
    ),
    (
        "颲脫平?冽?璆剖??芰靽蝵?,
        "?拙??乩??脫蝑ㄡ?啁??????靽風?芣??,
        "https://www.forest.gov.tw/",
    ),
    (
        "IUCN Red List",
        "??靽????澈嚗?亦???賊???鞎??????,
        "https://www.iucnredlist.org/",
    ),
    (
        "Cat Specialist Group",
        "鞎??????脰?銵?閮??閬???皞?,
        "https://catsg.org/",
    ),
    (
        "Taiwan Biodiversity Research Network",
        "?臭??箇蝔株??航??憭見?批辣隡豢閰Ｗ???,
        "https://taibif.tw/",
    ),
]


LITERATURE = [
    (
        "Vigne et al., 2016",
        "Earliest ?omestic??Cats in China Identified as Leopard Cat (Prionailurus bengalensis)",
        "PLOS ONE",
        "10.1371/journal.pone.0147295",
    ),
    (
        "Tamada et al., 2008",
        "Molecular Diversity and Phylogeography of the Asian Leopard Cat, Felis bengalensis, Inferred from Mitochondrial and Y-Chromosomal DNA Sequences",
        "Zoological Science",
        "10.2108/zsj.25.154",
    ),
    (
        "Chua et al., 2016",
        "Population density, spatiotemporal use and diet of the leopard cat (Prionailurus bengalensis) in a human-modified succession forest landscape of Singapore",
        "Mammal Research",
        "10.1007/s13364-015-0259-4",
    ),
    (
        "Grassman et al., 2005",
        "Spatial organization and diet of the leopard cat (Prionailurus bengalensis) in north-central Thailand",
        "Journal of Zoology",
        "10.1017/S095283690500659X",
    ),
    (
        "Johnson et al., 2006",
        "The Late Miocene Radiation of Modern Felidae: A Genetic Assessment",
        "Science",
        "10.1126/science.1122277",
    ),
]


def set_cell_text(cell, text: str, *, bold: bool = False, size_pt: float = 10.0) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size_pt)
    run.font.color.rgb = RGBColor(0, 0, 0)


def add_hyperlink(paragraph, text: str, url: str, *, size_pt: float = 10.0) -> None:
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_pr.append(r_fonts)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    r_pr.append(u)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size_pt * 2)))
    r_pr.append(sz)

    run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def apply_paragraph_style(paragraph, *, size_pt: float = 11, color: str = "000000",
                          bold: bool = False, before_pt: float = 0, after_pt: float = 6,
                          line_spacing: float = 1.25, alignment=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    paragraph.alignment = alignment
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before_pt)
    fmt.space_after = Pt(after_pt)
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = line_spacing
    for run in paragraph.runs:
        run.font.name = "Calibri"
        run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        run.font.size = Pt(size_pt)
        run.bold = bold
        run.font.color.rgb = RGBColor.from_string(color)


def style_heading(paragraph, level: int) -> None:
    if level == 1:
        size_pt, color, before_pt, after_pt = 16, "2E74B5", 18, 10
    elif level == 2:
        size_pt, color, before_pt, after_pt = 13, "2E74B5", 14, 7
    else:
        size_pt, color, before_pt, after_pt = 12, "1F4D78", 10, 5
    apply_paragraph_style(
        paragraph,
        size_pt=size_pt,
        color=color,
        bold=False,
        before_pt=before_pt,
        after_pt=after_pt,
        line_spacing=1.0,
    )


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def bold_first_paragraph(cell) -> None:
    for run in cell.paragraphs[0].runs:
        run.bold = True


def set_table_header(table, header_fill: str = "E8EEF5") -> None:
    for cell in table.rows[0].cells:
        shade_cell(cell, header_fill)
        set_cell_text(cell, cell.text, bold=True, size_pt=10.0)
        bold_first_paragraph(cell)


def build_document() -> Document:
    doc = Document()

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.25

    for name, size_pt, color, before_pt, after_pt in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size_pt)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before_pt)
        style.paragraph_format.space_after = Pt(after_pt)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(3)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = title.add_run("?唾?靽?亥?????)
    r.font.name = "Calibri"
    r._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    r._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    r.font.size = Pt(18)
    r.bold = True
    r.font.color.rgb = RGBColor(0, 0, 0)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(10)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sr = subtitle.add_run("?渡?摰蝬脤??誨銵冽抒?蝛嗆??鳴??拙?蝘??????鞈????其?皞?)
    sr.font.name = "Calibri"
    sr._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    sr._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    sr.font.size = Pt(10.5)
    sr.font.color.rgb = RGBColor(80, 80, 80)

    p = doc.add_paragraph()
    p.add_run("隤芣?嚗?).bold = True
    p.add_run("?遢鞈??芸??園??舐?仿???蝬脩??亙??DOI ???嚗迤撘???典?嚗遣霅啣??????詨?璅?????撟港遢??)
    apply_paragraph_style(p, size_pt=11, color="000000", bold=False, before_pt=0, after_pt=8, line_spacing=1.25)

    h1 = doc.add_paragraph("銝?雯??皞?)
    style_heading(h1, 1)

    web_intro = doc.add_paragraph()
    web_intro.add_run("撱箄降??摰???脩?蝜雯蝡霈嚗?鋆飛銵??颯?)
    apply_paragraph_style(web_intro, size_pt=11, color="000000", before_pt=0, after_pt=6, line_spacing=1.25)

    web_table = doc.add_table(rows=1, cols=3)
    web_table.style = "Table Grid"
    web_headers = ["鞈??迂", "?批捆??", "???"]
    for cell, text in zip(web_table.rows[0].cells, web_headers):
        cell.text = text
    for row in web_table.rows[:1]:
        for cell in row.cells:
            shade_cell(cell, "E8EEF5")
            set_cell_text(cell, cell.text, bold=True, size_pt=10.0)

    for name, note, url in WEB_RESOURCES:
        row = web_table.add_row()
        set_cell_text(row.cells[0], name, size_pt=10.0)
        set_cell_text(row.cells[1], note, size_pt=10.0)
        row.cells[2].text = ""
        para = row.cells[2].paragraphs[0]
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        add_hyperlink(para, "??蝬脩?", url, size_pt=10.0)

    apply_table_geometry(
        web_table,
        [2080, 3920, 3360],
        table_width_dxa=9360,
        indent_dxa=120,
    )

    h1b = doc.add_paragraph("鈭?蝛嗆???)
    style_heading(h1b, 1)

    lit_intro = doc.add_paragraph()
    lit_intro.add_run("隞乩?隞亥?蝘?憿?蝢斤????扯?瞍??賊??弦?箔蜓嚗靘踵??唾??曉?頛??渡?鞎??窗??)
    apply_paragraph_style(lit_intro, size_pt=11, color="000000", before_pt=0, after_pt=6, line_spacing=1.25)

    lit_table = doc.add_table(rows=1, cols=4)
    lit_table.style = "Table Grid"
    lit_headers = ["雿?撟港遢", "憿?", "??", "DOI"]
    for cell, text in zip(lit_table.rows[0].cells, lit_headers):
        cell.text = text
    for cell in lit_table.rows[0].cells:
        shade_cell(cell, "E8EEF5")
        set_cell_text(cell, cell.text, bold=True, size_pt=10.0)

    for authors, title_text, journal, doi in LITERATURE:
        row = lit_table.add_row()
        set_cell_text(row.cells[0], authors, size_pt=9.8)
        set_cell_text(row.cells[1], title_text, size_pt=9.8)
        set_cell_text(row.cells[2], journal, size_pt=9.8)
        row.cells[3].text = ""
        pdoi = row.cells[3].paragraphs[0]
        pdoi.paragraph_format.space_before = Pt(0)
        pdoi.paragraph_format.space_after = Pt(0)
        pdoi.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        add_hyperlink(pdoi, doi, f"https://doi.org/{doi}", size_pt=9.8)

    apply_table_geometry(
        lit_table,
        [1700, 4300, 2000, 1360],
        table_width_dxa=9360,
        indent_dxa=120,
    )

    h2 = doc.add_paragraph("銝?撅神雿?湔撘??暺?)
    style_heading(h2, 2)

    note = doc.add_paragraph()
    note.add_run("1. ")
    note.add_run("?唾?靽撣貉?銝駁??璉脣?渡???頝航楝畾箝?梯炊??鈭箇銵???)
    apply_paragraph_style(note, size_pt=11, color="000000", before_pt=0, after_pt=4, line_spacing=1.25)

    note2 = doc.add_paragraph()
    note2.add_run("2. ")
    note2.add_run("?亥?撘瑁矽?啁?典獢?嚗?芸??亙??嫣??脰????啁?唾?靽??嚗?撠摮貉????)
    apply_paragraph_style(note2, size_pt=11, color="000000", before_pt=0, after_pt=4, line_spacing=1.25)

    note3 = doc.add_paragraph()
    note3.add_run("3. ")
    note3.add_run("?亙??閬?憿??荔??航???惇?潸?蝘?????憿?銝西?鈭散?嗡??黎瘥???)
    apply_paragraph_style(note3, size_pt=11, color="000000", before_pt=0, after_pt=6, line_spacing=1.25)

    return doc


def main() -> None:
    doc = build_document()
    doc.save(OUT_PATH)
    print(f"saved {OUT_PATH.resolve()}")


if __name__ == "__main__":
    main()

