from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

PROJECT_ROOT = Path(__file__).resolve().parents[2]


OUT = PROJECT_ROOT / "AI?????_??????_??.docx"


def set_font(run, name="Calibri", size=11, bold=False, color="000000"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def para_spacing(p, before=0, after=0, line=1.15):
    fmt = p.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def table_borders(table, color="D9D9D9", size="8"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def style_doc(doc):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)

    for name, size, color in (("Heading 1", 16, "2E74B5"), ("Heading 2", 13, "2E74B5"), ("Heading 3", 12, "1F4D78")):
        st = doc.styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)


def h(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = doc.styles[f"Heading {level}"]
    p.paragraph_format.line_spacing = 1.1
    if level == 1:
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(8)
    elif level == 2:
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    else:
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_font(r, "Calibri", {1: 16, 2: 13, 3: 12}[level], False, {1: "2E74B5", 2: "2E74B5", 3: "1F4D78"}[level])
    return p


def body(doc, text, after=6):
    p = doc.add_paragraph()
    para_spacing(p, after=after)
    r = p.add_run(text)
    set_font(r, "Calibri", 11)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    para_spacing(p, after=4)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    r = p.add_run(text)
    set_font(r, "Calibri", 11)
    return p


def num(doc, text):
    p = doc.add_paragraph(style="List Number")
    para_spacing(p, after=4)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    r = p.add_run(text)
    set_font(r, "Calibri", 11)
    return p


doc = Document()
style_doc(doc)

# Title
p = doc.add_paragraph()
para_spacing(p, after=3, line=1.0)
r = p.add_run("AI 撠??弦?∴??嗆?瘛箏控蝎暸?")
set_font(r, "Calibri", 26)

p = doc.add_paragraph()
para_spacing(p, after=8, line=1.0)
r = p.add_run("?啁?唾?靽 ? Google AI ??賣???冽?獢?)
set_font(r, "Calibri", 11, False, "555555")

note = doc.add_table(rows=1, cols=1)
note.alignment = WD_TABLE_ALIGNMENT.LEFT
note.autofit = False
note.columns[0].width = Inches(6.5)
table_borders(note, "D9D9D9", "8")
cell = note.cell(0, 0)
cell_margins(cell, 120, 120, 120, 120)
shade(cell, "F4F6F9")
p = cell.paragraphs[0]
para_spacing(p, after=0)
r = p.add_run("?祆?隞嗆?玨蝔之蝬晞?蝔玨蝔???閰喟敦??嚗?冽 Chromebook 撖虫?隤脰?撅?隤脯?)
set_font(r, "Calibri", 11)

h(doc, "銝?玨蝔蜇閬?, 1)
overview = doc.add_table(rows=7, cols=2)
overview.alignment = WD_TABLE_ALIGNMENT.LEFT
overview.autofit = False
overview.columns[0].width = Inches(1.7)
overview.columns[1].width = Inches(4.8)
table_borders(overview, "D9D9D9", "8")

overview_items = [
    ("隤脩??迂", "AI 撠??弦?∴??嗆?瘛箏控蝎暸?"),
    ("銝駁??詨?", "隞亦???脩銝駁?嚗???Teachable Machine?oogle AI Studio / Gemini?otebookLM??),
    ("?拍撠情", "??擃僑蝝?葉雿僑蝝?),
    ("撱箄降??", "80 ??嚗? 蝭隤莎??亦撅??臬?蝮桃 40 ??50 ??"),
    ("摮貊???", "?唾????I 敶勗?????憭??豢????霅澈?Ｙ弦?恐撠???"),
    ("?飛?孵?", "雓膩?內蝭祕雿?蝯?雿???隢?),
    ("閮剖?", "Chromebook?雯頝胯oogle 撣唾??eachable Machine?oogle AI Studio?otebookLM??敶梯身??),
]
for i, (k, v) in enumerate(overview_items):
    for j, txt in enumerate((k, v)):
        c = overview.cell(i, j)
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell_margins(c, 100, 120, 100, 120)
        if j == 0:
            shade(c, "E8EEF5")
        p = c.paragraphs[0]
        para_spacing(p, after=0)
        r = p.add_run(txt)
        set_font(r, "Calibri", 11, j == 0)

h(doc, "鈭?摮詨摰孵之蝬?, 1)
for t in [
    "隤脩?撠嚗?霅??瘛箏控??嚗?閫??憭??豢??矽?亦??,
    "璅∠?銝嚗蝙??Teachable Machine 閮毀?餉閬死???具?,
    "璅∠?鈭??拍 Google AI Studio / Gemini ???豢??抒?嚗芋?砍?撠??飛摰嗚?,
    "璅∠?銝??拍 NotebookLM 銝靽鞈?嚗脰?銝駁????霅??,
    "璅∠???雿輻 NotebookLM ???唾??飛蝪∪???脰??怨?閮?銵具?,
]:
    bullet(doc, t)

h(doc, "銝玨蝔???, 1)
for t in [
    "蝚砌?蝭???曉 AI 敶勗?颲刻?閮毀?蝷?敹萄遣蝡?,
    "蝚砌?蝭???曉蝝?蝺璈???otebookLM ?亥?摨急蝛嗉?摰??????,
    "?交撅?隤莎??臭?閮剖??飛??摨西矽?渡蝎曄陛??蝔?雿????芋蝯?摰?摩??,
]:
    bullet(doc, t)

h(doc, "?玨蝔?蝔?, 1)
flow = doc.add_table(rows=1, cols=4)
flow.alignment = WD_TABLE_ALIGNMENT.LEFT
flow.autofit = False
for idx, width in enumerate([Inches(0.9), Inches(2.1), Inches(1.8), Inches(1.7)]):
    flow.columns[idx].width = width
table_borders(flow, "D9D9D9", "8")
for i, head in enumerate(["??", "?飛?批捆", "?葦瘣餃?", "摮貊?隞餃?"]):
    c = flow.rows[0].cells[i]
    shade(c, "E8EEF5")
    cell_margins(c)
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(head)
    set_font(r, "Calibri", 11, True)

rows = [
    ("0-8 ??, "撠??憓遣蝡?, "?箇內?唾???憭??豢??抒?嚗???韏瑁?頞??, "閫撖?皜研?蝑?),
    ("8-23 ??, "Teachable Machine 撖虫?", "蝷箇?撱箇??????喳???蝺渲?皜祈岫??, "頝?摰?璅∪?閮毀?葫閰艾?),
    ("23-30 ??, "璅∪?閮?", "撘?摮貊???隤文????脫撘?, "?澈蝯??耨甇?瘜?),
    ("30-45 ??, "Gemini 敶勗???", "蝷箇?雿輻 Google AI Studio / Gemini ???豢??抒???, "閫撖撓?箝?撠?摰嫘?),
    ("45-57 ??, "NotebookLM ??", "銝?唾?靽鞈?嚗內蝭????亥??渡???, "?曉?????獢?),
    ("57-70 ??, "摰??????, "蝷箇???蝪∪??閮?銵具?, "閫撖????箔耨甇??),
    ("70-80 ??, "蝮賜??潸”", "?“?芋蝯?撣嗅飛??剔銵具?, "?澈摮貊?????),
]
for row in rows:
    cells = flow.add_row().cells
    for i, txt in enumerate(row):
        c = cells[i]
        cell_margins(c, 100, 120, 100, 120)
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i else WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(txt)
        set_font(r, "Calibri", 10.5, i == 1)

h(doc, "鈭底蝝唳?摮詨摰?, 1)
h(doc, "1. 隤脩?撠", 2)
for t in [
    "?箇內?唾???憭??豢??抒?嚗?摮貊??牧?箇??啁???憓?蝝Ｕ?,
    "撘?摮貊???憒??抒?敺?嚗??獐敹恍??AI ?臭誑?獐撟怠?嚗?,
]:
    bullet(doc, t)

h(doc, "2. 璅∠?銝嚗??儘霅芋??蝺?, 2)
for t in [
    "雿輻 Teachable Machine 撱箇??餉閬死???具?,
    "撱箄降憿?航身?箇?振鞎隞??抬???隤脣?蝝?隤踵??,
    "霈飛??閫??AI 銝憭拍?????????飛蝧?,
    "??閬雲憭???摨西?憭見???航??⊿??嚗芋??蝛拙???,
]:
    bullet(doc, t)

h(doc, "3. 璅∠?鈭?璅⊥撠???摮詨振", 2)
for t in [
    "撠?憭??豢??抒?銝??Google AI Studio / Gemini??,
    "雿輻?箏??內閰? AI 颲刻??蝔桅?????箄??航憡???,
    "撘?摮貊?瘥? AI ?方??撌梯?撖?虫??氬?,
    "??摮貊?嚗I ?臬?甇亙????銝?誨鈭箏極蝣箄???,
]:
    bullet(doc, t)

h(doc, "4. 璅∠?銝?NotebookLM ?亥?摨急蝛?, 2)
for t in [
    "銝?唾?靽鞈????格?蝡撠?摰????,
    "??銝駁????渡??唾?璉脣????靽銵???,
    "撱箇?隤脣??亥?摨恬?霈飛??????皜?嚗?蝑?皞Ⅱ??,
]:
    bullet(doc, t)

h(doc, "5. 璅∠????芸???摰????, 2)
for t in [
    "雿輻 NotebookLM ???唾??飛蝪∪??,
    "雿輻 NotebookLM ??靽閮鞈??”??,
    "撠霅???澈?撅內??∪?摰??????,
]:
    bullet(doc, t)

h(doc, "?准底蝝唳?獢?, 1)
h(doc, "?飛?格?", 2)
for t in [
    "隤??唾?????靽???扼?,
    "摮豢?雿輻 Teachable Machine ?脰? AI 敶勗?????,
    "?圾 Google AI Studio / Gemini ?舐?潛????璉脣?方???,
    "摮豢?雿輻 NotebookLM ?脰?鞈??渡???蝑?摰??????,
]:
    bullet(doc, t)

h(doc, "?飛??", 2)
for t in [
    "?唾??臬??滓撅梢?閬????押?,
    "蝝?蝺璈??隤踵??閬極?瑯?,
    "AI ?臭誑撟怠?????閬?雿?蝯霈隞?鈭粹?蝣箄???,
    "鞈??渡??質???蝪∪??閮?銵刻??∪?摰???批捆??,
]:
    bullet(doc, t)

h(doc, "?葦雓阮", 2)
script = [
    ("?", "隞予????AI 撠??弦?∴?銝韏瑕鼠?唾????脖遙??),
    ("?唾?閫撖?, "?唾??絲靘?鞎?雿??臬???閬????嚗?瘣餃瘛箏控?啣???),
    ("AI ?圾", "AI 銝?芸楛?仿?蝑?嚗???飛蝧?鞈?頞末嚗????舫???),
    ("??撘?", "????蝺湔芋??? Gemini ????敺 NotebookLM ?霅??摰????),
    ("?嗆?", "隞予摮詨???芣??撌亙嚗??????閫??嗚?敹憓?銵???),
]
for title, text in script:
    p = doc.add_paragraph()
    para_spacing(p, before=2, after=4)
    r1 = p.add_run(title + "嚗?)
    set_font(r1, "Calibri", 11, True)
    r2 = p.add_run(text)
    set_font(r2, "Calibri", 11)

h(doc, "銝?湔雿輻??蝷箄?", 1)
prompt_box = doc.add_table(rows=3, cols=1)
prompt_box.alignment = WD_TABLE_ALIGNMENT.LEFT
prompt_box.autofit = False
prompt_box.columns[0].width = Inches(6.5)
table_borders(prompt_box, "D9D9D9", "8")
prompt_texts = [
    "Teachable Machine嚗?靘玨???遣蝡??/ 摰嗉? / ?嗡????憿芋??銝西?撖芋??血????摨西炊?扎?,
    "Gemini ??嚗??臭?雿?撠??飛摰嗅?????撐蝝?蝺璈??頛詨??迂????箝憓???璉脣?孵?撱箄降??,
    "NotebookLM ??嚗??寞????喟??唾?靽鞈?嚗??ㄡ?啜蜓閬???瘞?舀??靽銵?嚗蒂?冽???????,
]
for i, txt in enumerate(prompt_texts):
    c = prompt_box.cell(i, 0)
    cell_margins(c, 120, 120, 120, 120)
    if i == 0:
        shade(c, "F4F6F9")
    p = c.paragraphs[0]
    para_spacing(p, after=0)
    r = p.add_run(txt)
    set_font(r, "Calibri", 11)

h(doc, "?怒??撘?, 1)
for t in [
    "????嚗隤芸?唾??ㄡ?啜???靽銵???,
    "撖虫?閫撖??賢??撠???AI 撌亙??雿?蝔?,
    "鞈??方?嚗颲典 AI 頛詨?臬????,
    "???Ｗ嚗摰?銝隞賜陛?晞?閬?鞈??”????,
]:
    bullet(doc, t)

h(doc, "銋?摮豢???, 1)
for t in [
    "憒? NotebookLM ?撠??蝪∪??閮?銵刻撓?綽??舀隞交?閬???暺??葦蝷箇????蹂誨??,
    "Gemini / AI Studio ?????Ⅱ璅內?箝?甇亙???銝?湔?嗆?蝯?獢?,
    "撱箄降隤脣????末???????踹?隤脣???鋡怎???撠???,
]:
    bullet(doc, t)

body(doc, "?辣蝯?嚗???臭??箏?瞍玨?楊?玨蝔?鞈????嗥?摮豢?玨?敹???)

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)


